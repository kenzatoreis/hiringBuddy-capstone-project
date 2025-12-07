# -----------------------------------------------
# HiringBuddy AI Core _ all agent logic and persistent with (pgvector)
# -----------------------------------------------

import io, json, re, time
from typing import List, Dict, Optional, Any
from fastapi import APIRouter, UploadFile, File, HTTPException, Form, Body, Depends, Request
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PyPDF2 import PdfReader
from botocore.exceptions import ClientError
from sqlalchemy.orm import Session
from bedrock_client import invoke_chat, embed_text, DEFAULT_AWS_REGION, DEFAULT_CLAUDE
from db import get_db
from models import Resume, ResumeChunk, MatchAttempt, InterviewAttempt
from auth import get_user_from_token
from fastapi.responses import StreamingResponse
import re
import os
import httpx
import pdfplumber
from dotenv import load_dotenv
from pydantic import BaseModel
import random
router = APIRouter(prefix="/ai", tags=["ai"])
load_dotenv()  
SERPER_API_KEY = os.getenv("SERPER_API_KEY", "")
SERPER_ENDPOINT = "https://google.serper.dev/search"
#for lang options english and french 
def _lang_hint(language: str) -> str:
    lang = (language or "en").lower()
    if lang == "fr":
        return (
            "IMPORTANT: Rédige toutes les réponses (highlights, missing, questions, feedback, etc.) "
            "en français naturel, clair et professionnel."
        )
    return (
        "IMPORTANT: Write all responses (highlights, missing, questions, feedback, etc.) "
        "in clear, natural English."
    )

def token_aware_chunks(text: str, max_tokens: int = 700, overlap: int = 80):
    import re, math
    text = re.sub(r'\s+', ' ', text).strip()
    sentences = re.split(r'(?<=[.!?])\s+', text)
    chunks, cur, cur_tokens = [], [], 0

    def est_tokens(s: str) -> int:
        return max(1, int(len(s.split()) / 0.75))

    i = 0
    while i < len(sentences):
        s = sentences[i]
        t = est_tokens(s)
        if cur_tokens + t <= max_tokens or not cur:
            cur.append(s); cur_tokens += t; i += 1
        else:
            chunk = ' '.join(cur).strip()
            if chunk:
                chunks.append(chunk)
            prefix, tok = [], 0
            for s_back in reversed(cur):
                tok += est_tokens(s_back)
                prefix.append(s_back)
                if tok >= overlap:
                    break
            cur = list(reversed(prefix))
            cur_tokens = sum(est_tokens(x) for x in cur)
    last = ' '.join(cur).strip()
    if last:
        chunks.append(last)
    return chunks


# ---------------- claude Wrapper ----------------

def safe_invoke_chat(*args, retries=2, delay=4, **kwargs):
    last_err = None

    for attempt in range(1, retries + 1):
        try:
            return invoke_chat(*args, **kwargs)
        except ClientError as e:
            code = e.response.get("Error", {}).get("Code", "")
            # serve rdown retry
            if code in ("ThrottlingException", "ServiceUnavailableException"):
                last_err = e
                time.sleep(delay * attempt)
                continue
            
            raise
        except Exception as e:
            msg = str(e)
            # retry on timeouts too
            if "ReadTimeoutError" in msg or "timed out" in msg.lower():
                last_err = e
                if attempt < retries:
                    time.sleep(delay * attempt)
                    continue
                # after last attempt, return error
                raise HTTPException(504, "Claude timeout. Try again later.")
        
            raise

    # all retries exhausted
    raise HTTPException(504, "Claude overloaded, retry later.")


# ---------------- File Text Extraction ----------------
def _extract_text_from_docx_bytes(b: bytes) -> str:
    with io.BytesIO(b) as buf:
        doc = Document(buf)
        return "\n".join(p.text for p in doc.paragraphs)

def _extract_text_from_pdf_bytes(b: bytes) -> str:
    # reader = PdfReader(io.BytesIO(b)) 
    # return "\n".join((p.extract_text() or "") for p in reader.pages)
    """
    try pdfplumber first 
    """
    # --- pdfplumber <3<3
    try:
        text_chunks = []
        with io.BytesIO(b) as buf:
            with pdfplumber.open(buf) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text() or ""
                    text_chunks.append(page_text)

        txt = "\n".join(text_chunks)
        #cleanup
        txt = re.sub(r"[ \t]+", " ", txt)
        txt = re.sub(r"\n{3,}", "\n\n", txt)
        txt = txt.strip()

        if txt:
            return txt
    except Exception as e:
        print("[WARN] pdfplumber failed, falling back to PyPDF2:", e)

    # ---  PyPDF2 hna as fallback
    reader = PdfReader(io.BytesIO(b))
    txt = "\n".join((p.extract_text() or "") for p in reader.pages)
    txt = re.sub(r"[ \t]+", " ", txt)
    txt = re.sub(r"\n{3,}", "\n\n", txt)
    return txt.strip()


def _extract_text(file: UploadFile, raw: bytes) -> str:
    name = (file.filename or "").lower()
    if name.endswith(".docx"):
        return _extract_text_from_docx_bytes(raw)
    if name.endswith(".pdf"):
        return _extract_text_from_pdf_bytes(raw)
    if name.endswith(".txt"):
        return raw.decode("utf-8", errors="ignore")
    raise HTTPException(400, "Unsupported file type (.docx, .pdf, .txt only)")

# ---------------- cosine Ssimilar-----------
def cos_sim(a: List[float], b: List[float]) -> float:
    if not a or not b or len(a) != len(b):
        return 0.0
    import math
    dot = sum(x*y for x, y in zip(a, b))
    norm_a = math.sqrt(sum(x*x for x in a))
    norm_b = math.sqrt(sum(x*x for x in b))
    if norm_a == 0 or norm_b == 0:
        return 0.0
    return dot / (norm_a * norm_b)

# ---------------- /peek_doc ----------------
@router.post("/peek_doc")
async def peek_doc(file: UploadFile = File(...), limit: int = 1200):
    raw = await file.read()
    if not raw:
        raise HTTPException(400, "Empty file.")
    text = _extract_text(file, raw)
    if not text.strip():
        raise HTTPException(400, "No readable text.")
    return {"chars": len(text), "head": text[:max(100, min(limit, 4000))], "full": text}

# ---------------- /index_resume_mem ----------------
@router.post("/index_resume_mem")
async def index_resume_mem(
    request: Request,
    file: UploadFile = File(...),
    candidate_name: str = Form("Unknown"),
    db: Session = Depends(get_db),
):
    user = get_user_from_token(request, db)
    raw = await file.read()
    if not raw:
        raise HTTPException(400, "Empty file.")
    text = _extract_text(file, raw)
    if not text.strip():
        raise HTTPException(400, "No readable text.")

    # save resume record
    resume = Resume(user_id=user.id, name=file.filename, text=text)
    db.add(resume)
    db.commit()
    db.refresh(resume)

    # split and embed each para
    paragraphs = token_aware_chunks(text, max_tokens=700, overlap=80)
    if len(paragraphs) <= 1:
        words = text.split()
        paragraphs = [" ".join(words[i:i+800]) for i in range(0, len(words), 700)]

    if not paragraphs:
        paragraphs = [text]

    for p in paragraphs:
        try:
            emb = embed_text(p, aws_region=DEFAULT_AWS_REGION)
            if not emb or not isinstance(emb, list):
                raise ValueError("Empty embedding response")
        except Exception as e:
            print(f"[WARN] Embedding failed for chunk: {p[:60]!r} ({e})")
            # fallback vector 256 dims)
            emb = [0.0] * 256

        db.add(ResumeChunk(resume_id=resume.id, text=p, embedding=emb))

    db.commit()
    print(f"[OK] Indexed resume #{resume.id} with {len(paragraphs)} chunks")
    return {"ok": True, "resume_id": resume.id}


# ---------------- Helper: Retrieve top snippets ----------------
def best_snippets_from_db(requirement: str, db, user, region=DEFAULT_AWS_REGION,
                          top_k_resumes=2, top_k_snippets=3):
    qv = embed_text(requirement, aws_region=region)
    keywords = [w for w in re.findall(r"\w+", requirement.lower()) if len(w) > 2]
    scored = []

    for resume in (
    db.query(Resume)
      .filter(Resume.user_id == user.id)
      .order_by(Resume.created_at.desc())
      .limit(1)
      .all()
):
        chunk_scores = []
        for ch in resume.chunks:
            # convert JSON from db to floats
            vec = []
            for x in (ch.embedding or []):
                try:
                    vec.append(float(x))
                except Exception:
                    continue

            sem = cos_sim(qv, vec)
            kw_hits = sum(1 for w in keywords if w in ch.text.lower())
            score = 0.85 * sem + 0.15 * (min(kw_hits, 5) / 5.0)
            chunk_scores.append((score, ch.text))

        if not chunk_scores:
            continue
        chunk_scores.sort(reverse=True, key=lambda t: t[0])
        best = chunk_scores[0][0]
        top = [t[1] for t in chunk_scores[:top_k_snippets]]
        # --- append Sills + Experience blocks k---
        full_text = resume.text or ""

        skills_block = _extract_section_block(full_text, ["SKILLS", "TECHNICAL SKILLS"])
        experience_block = _extract_section_block(full_text, ["EXPERIENCE", "PROFESSIONAL EXPERIENCE"])

        # If found, append them to snippets shown to Claude
        if skills_block:
            top.append(skills_block)

        if experience_block:
            top.append(experience_block)

        print(f"[{resume.name}] best semantic score = {best:.3f}")
        scored.append((best, resume, top))

    scored.sort(reverse=True, key=lambda t: t[0])
    return scored[:top_k_resumes]


# ---------------- /match_mem/ matching cv jd----------------
@router.post("/match_mem")
def match_mem(
    request: Request,
    requirement: str = Body(..., embed=True),
    language: str = Body("en"),
    db: Session = Depends(get_db),
):
    user = get_user_from_token(request, db)
    if not requirement.strip():
        raise HTTPException(400, "Empty requirement.")
    language = (language or "en").lower()
    if language not in {"en", "fr"}:
        language = "en"

    scored = best_snippets_from_db(requirement, db, user)
    results = []

    for best, resume, snippets in scored:
        lang_hint = _lang_hint(language)

        prompt = f"""{lang_hint}
    You are a hiring assistant. Return ONLY JSON: {{
  "score": 0-100,
  "highlights": [string],
  "missing": [string]
}}
Requirement:
---
{requirement}
---
Top evidence excerpts (from the candidate):
---
{chr(10).join(snippets)}
---"""

        out = safe_invoke_chat(
            messages=[{"role": "user", "content": [{"text": prompt}]}],
            system=[{"text": "Return ONLY valid JSON. No prose."}],
            model_id=DEFAULT_CLAUDE,
            aws_region=DEFAULT_AWS_REGION,
            max_tokens=380,
            temperature=0.0,
        )
        cleaned = out.strip().strip("`").replace("json", "", 1).strip()
        # print claude returned
        print("\n================ CLAUDE RAW OUTPUT ================")
        print(cleaned)
        print("===================================================\n")
        results.append({
            "resume_id": resume.id,
            "candidate": resume.name,
            "retrieval_semantic_best": best,
            "llm_json": cleaned
        })
        time.sleep(1.5)
#             # existing loop:
#     for best, resume, snippets in scored:
#         prompt = f"""You are a hiring assistant. Return ONLY JSON: {{
#   "score": 0-100,
#   "highlights": [string],
#   "missing": [string]
# }}
# Requirement:
# ---
# {requirement}
# ---
# Top evidence excerpts (from the candidate):
# ---
# {chr(10).join(snippets)}
# ---"""

#         out = safe_invoke_chat(
#             messages=[{"role": "user", "content": [{"text": prompt}]}],
#             system=[{"text": "Return ONLY valid JSON. No prose."}],
#             model_id=DEFAULT_CLAUDE,
#             aws_region=DEFAULT_AWS_REGION,
#             max_tokens=380,
#             temperature=0.0,
#         )
#         cleaned = out.strip().strip("`").replace("json", "", 1).strip()
#         print("\n================ CLAUDE RAW OUTPUT ================")
#         print(cleaned)
#         print("===================================================\n")

#         results.append({
#             "resume_id": resume.id,
#             "candidate": resume.name,
#             "retrieval_semantic_best": best,
#             "llm_json": cleaned
#         })
#         time.sleep(1.5)

    # ---- NEW: log history for the *best* result ----
    best_result = max(results, key=lambda r: r.get("retrieval_semantic_best", 0), default=None)

    if best_result is not None:
        try:
            best_json = best_result["llm_json"]
            data = json.loads(best_json) if isinstance(best_json, str) else best_json
            raw_score = data.get("score", 0)
            score_val = int(max(0, min(100, round(float(raw_score)))))
        except Exception:
            score_val = None


        # short label for the JD
        job_title = (requirement or "").strip().split("\n")[0][:255]
        jd_snippet = (requirement or "")[:1000]

        attempt = MatchAttempt(
            user_id=user.id,
            resume_id=best_result.get("resume_id"),
            job_title=job_title,
            score=score_val,
            jd_snippet=jd_snippet,
            raw_json=best_result["llm_json"],
        )
        db.add(attempt)
        db.commit()

    return {"results": results}
#----ats system----
class ExtractRequest(BaseModel):
    text: str

def clean_json_generic(s: str):
    s = s.strip()
    s = s.replace("```json", "").replace("```", "")
    return s
#---ats
@router.post("/extract_keywords")
def extract_keywords(body: ExtractRequest):
    """
    Pure keyword extraction — NOT semantic, NOT CS-specific.
    Extracts concrete requirement tokens from ANY job description.
    """

    prompt = f"""
Extract the DISTINCT keywords from this job description.
RULES:
- Include tools, methods, certifications, frameworks, techniques,
  software names, domain-specific terms, methodologies,
  hard skills, platforms, and technical nouns.
- Keep meaningful 1–3 word phrases.
- NO soft skills (communication, teamwork, leadership, etc.)
- Output ONLY a JSON array of strings.
Text:
{body.text}
"""

    out = safe_invoke_chat(
        messages=[{"role": "user", "content": [{"text": prompt}]}],
        system=[{"text": "Return ONLY a JSON array. No prose. No explanations."}],
        model_id=DEFAULT_CLAUDE,
        aws_region=DEFAULT_AWS_REGION,
        temperature=0.0,
        max_tokens=200,
    )

    cleaned = clean_json_generic(out)

    try:
        arr = json.loads(cleaned)
        if isinstance(arr, list):
            return {"keywords": arr}
    except:
        pass

    return {"keywords": []}


#-----suggest-----
@router.post("/suggestions_for_improvement")
def suggestions_for_improvement(
    resume_text: str = Body(...),
    job_description: str = Body(...),
    missing: List[str] = Body(default=[]),
):
    """
    Generate realistic improvement suggestions (certificates, projects, skills to learn)
    without modifying the resume content.
    """
    if not resume_text or not job_description:
        raise HTTPException(400, "resume_text and job_description are required.")

    prompt = f"""
You are a career advisor for a Computer Science student with an AI / software focus.

Your job is to suggest **interesting, specific and realistic** ways to improve the profile
for the given job description.

VERY IMPORTANT RULES
- Do NOT invent fake degrees, companies or long full-time experiences.
- You MAY suggest realistic online certificates / MOOCs / nanodegrees (Coursera, edX, Udemy, AWS, Google, etc.).
- You MAY suggest side-projects, GitHub projects, hackathons, Kaggle competitions, etc.
- ALWAYS keep it tailored to Computer Science / software / data, not random fields.
- Write everything in **clear, natural English** (no French, no mixed languages).

Focus on:
1) **Certificates / trainings** to close the most important gaps.
   - Name concrete programs when possible (e.g. "AWS Certified Developer – Associate",
     "Spring Framework & Hibernate bootcamp", "Docker & Kubernetes for Java Developers").
   - Prefer 4–6 items that clearly match the JD technologies or missing skills.

2) **Practical project ideas** (side-projects) that the student could add to their CV.
   - Each project should be 1 short line, mention stack + goal
     (e.g. "Build a RESTful API with Spring Boot + PostgreSQL for managing ...").
   - Prefer 4–6 concrete ideas, not generic "do more projects".

3) **Skills / tools to learn**.
   - Focus on frameworks, libraries, cloud platforms, and dev practices that are
     clearly relevant for the job description and missing skills.
   - Prefer 6–10 short skill tokens (e.g. "Spring Security", "Docker Compose", "CI/CD with GitHub Actions").

Also include a short motivational note that feels personal and encouraging for
a final-year CS student / fresh graduate.

INPUTS
-------
RESUME:
{resume_text[:10000]}

JOB DESCRIPTION:
{job_description[:10000]}

MISSING SKILLS (from previous match step, can be empty):
{', '.join(missing) if missing else 'none'}

Return ONLY valid JSON (no markdown, no backticks) with this shape:
{{
  "certificates": [string],     // at least 4 items if possible
  "projects": [string],         // at least 4 items if possible
  "skills_to_learn": [string],  // 6–10 short tokens
  "note": "one or two sentences of motivation in English"
}}
"""

    out = safe_invoke_chat(
        messages=[{"role": "user", "content": [{"text": prompt}]}],
        system=[{"text": "Return ONLY valid JSON. No prose. No backticks."}],
        model_id=DEFAULT_CLAUDE,
        aws_region=DEFAULT_AWS_REGION,
        max_tokens=700,
        temperature=0.7,  # a bit more creative than before
    )

    cleaned = out.strip().strip("`").replace("json", "", 1).strip()
    try:
        data = json.loads(cleaned)
    except Exception:
        data = {"raw": cleaned}

    # safety: ensure keys exist so frontend never explodes
    if isinstance(data, dict):
        data.setdefault("certificates", [])
        data.setdefault("projects", [])
        data.setdefault("skills_to_learn", [])
        data.setdefault("note", "")

    return {"ok": True, "suggestions": data}



def _pick_resume_text(user_text: str, db_text: str) -> str:
    ut = (user_text or "").strip()
    dt = (db_text or "").strip()
    return ut if len(ut) >= len(dt) else dt
SECTION_HEADERS = {
    "PROFILE",
    "PROFESSIONAL EXPERIENCE", "EXPERIENCE",
    "PROJECTS PORTFOLIO", "PROJECTS",
    "EDUCATION",
    "CERTIFICATES",
    "SKILLS", "TECHNICAL SKILLS",
    "LANGUAGES",
}

def _canon(h: str) -> str:
    return re.sub(r"[^A-Z ]", "", (h or "").upper()).strip()

def _extract_section_block(full_text: str, start_labels: List[str]) -> str:
    t = (full_text or "").replace("\r", "\n")
    t = re.sub(r"[ \t]+", " ", t)
    lines = [ln.strip() for ln in t.splitlines() if ln.strip()]

    heads = []
    for i, ln in enumerate(lines):
        up = _canon(ln)
        if up in SECTION_HEADERS:
            heads.append((i, up))

    wanted = {_canon(s) for s in start_labels}
    start_idx = None
    for i, up in heads:
        if up in wanted:
            start_idx = i
            break
    if start_idx is None:
        return ""

    end_idx = len(lines)
    for i, up in heads:
        if i > start_idx:
            end_idx = i
            break

    block = "\n".join(lines[start_idx + 1:end_idx]).strip()
    block = re.sub(r"\n{3,}", "\n\n", block)
    return block
#helper for serper 
def _extract_skill_tokens_from_resume(text: str, max_skills: int = 12) -> List[str]:
    """
    Very lightweight skill extractor for building a Serper search query.
    Uses the SKILLS / TECHNICAL SKILLS sections if present.
    """
    skills_block = _extract_section_block(text, ["SKILLS", "TECHNICAL SKILLS"])
    if not skills_block:
        # Fallback with 800 tokens
        skills_block = text[:800]

    # Split on commas / semicolons / bullets
    raw_tokens = re.split(r"[,\n;/•·\-]+", skills_block)
    tokens = []
    for tok in raw_tokens:
        t = tok.strip()
        if len(t) < 2:
            continue
        #emove very generic words
        if t.lower() in {"skills", "tools", "languages", "projects", "experience"}:
            continue
        tokens.append(t)

    # deduplicate while preserving order.........
    seen = set()
    uniq = []
    for t in tokens:
        low = t.lower()
        if low in seen:
            continue
        seen.add(low)
        uniq.append(t)

    return uniq[:max_skills]

def _extract_job_profile_from_resume(text: str, max_skills: int = 20) -> dict:
    """
    Use Claude to extract a compact job profile from the resume:
    - technical skills
    - likely job titles
    - seniority level / years of experience
    """
    import json, re

    prompt = f"""
You are a career assistant.

Analyze the following resume text and extract information ONLY if it is clearly stated.

Return STRICT JSON with this schema:
{{
  "skills": [string],              // technical tools, languages, frameworks, cloud, DBs, etc. No dates, no soft-skill sentences.
  "primary_roles": [string],       // typical job titles this person could apply for (e.g. "Software Engineer", "Data Analyst").
  "seniority_level": "student" | "entry-level" | "junior" | "mid" | "senior",
  "years_of_experience": "0" | "0-1" | "1-2" | "2-3" | "3-5" | "5+"
}}

Rules:
- "skills": prefer compact tokens like "Python", "Django", "PostgreSQL", "React".
  Remove generic phrases like "Tools & Programming Languages", months/years, or committee roles.
- Only include skills that clearly look like technologies / tools / languages / libraries.
- If the candidate is still a student with only internships/projects, use
  seniority_level="student" and years_of_experience="0-1".
- If experience is not clear, be conservative (student/entry-level).

RESUME:
{text[:12000]}
"""

    out = safe_invoke_chat(
        messages=[{"role": "user", "content": [{"text": prompt}]}],
        system=[{"text": "Return ONLY valid JSON. No prose. No backticks."}],
        model_id=DEFAULT_CLAUDE,
        aws_region=DEFAULT_AWS_REGION,
        max_tokens=700,
        temperature=0.1,
    )

    raw = out.strip().strip("`").replace("json", "", 1).strip()
    try:
        profile = json.loads(raw)
    except Exception:
        profile = {}

    skills = profile.get("skills") or []
    if not isinstance(skills, list):
        skills = []

    # Cleanup: remove dates, long sentences, etc.
    cleaned = []
    for s in skills:
        if not isinstance(s, str):
            continue
        t = s.strip()
        if not t:
            continue
        if len(t) > 40:
            continue
        # exclude obvious dates/years
        if re.search(r"\b(20\d{2}|19\d{2})\b", t):
            continue
        # exclude generic labels
        if t.lower() in {
            "tools", "tools & programming languages", "skills", "soft skills",
            "methodologies & concepts", "methodologies", "concepts"
        }:
            continue
        cleaned.append(t)

    # de-dupe
    seen = set()
    uniq = []
    for t in cleaned:
        low = t.lower()
        if low in seen:
            continue
        seen.add(low)
        uniq.append(t)

    profile["skills"] = uniq[:max_skills]

    if not isinstance(profile.get("primary_roles"), list):
        profile["primary_roles"] = []

    profile.setdefault("seniority_level", "student")
    profile.setdefault("years_of_experience", "0-1")

    return profile

# ----draft cv -----
@router.post("/draft_cv_with_headers")
def draft_cv_with_headers(
    request: Request,
    resume_text: str = Body(""),
    job_description: str = Body(...),
    missing: List[str] = Body(default=[]),
    headers: List[dict] = Body(default=[
        {"title": "Profile", "context": "(keep it brief)"},
        {"title": "Professional Experience", "context": ""},
        {"title": "Education", "context": ""},
        {"title": "Projects", "context": ""},
        {"title": "Certificates", "context": ""},
        {"title": "Skills", "context": ""},
        {"title": "Languages", "context": ""},
    ]),
    language: str = Body("en"),
    db: Session = Depends(get_db),
):
    if not job_description:
        raise HTTPException(400, "job_description is required.")
    language = (language or "en").lower()
    if language == "fr":
        lang_hint = "Rédige toutes les sections du CV en français professionnel, clair, adapté à un étudiant AUI niveau junior."
    else:
        lang_hint = "Write all CV sections in clear, professional English, suitable for an AUI junior student."
    # ----- choose the longest resume text  -----
    user = get_user_from_token(request, db)
    ui_text = (resume_text or "").strip()
    db_text = ""
    latest = (
        db.query(Resume)
          .filter(Resume.user_id == user.id)
          .order_by(Resume.created_at.desc())
          .first()
    )
    if latest and (latest.text or "").strip():
        db_text = latest.text

    use_text = _pick_resume_text(ui_text, db_text)
    if not use_text:
        raise HTTPException(400, "No resume text found. Upload & index a CV, or paste resume_text.")

    # ----- anchor key blocks for fallback -----
    exp_block  = _extract_section_block(use_text, ["PROFESSIONAL EXPERIENCE", "EXPERIENCE"])
    proj_block = _extract_section_block(use_text, ["PROJECTS PORTFOLIO", "PROJECTS"])

    header_text = "\n".join([f"### {h['title']}: {h.get('context','')}" for h in headers])

    prompt = f""" 
You are an expert CV drafter. Adapt the candidate’s resume to the job description WITHOUT losing facts.
{lang_hint}
Strict rules:
1) DO NOT invent anything not explicitly present in the resume.
2) PRESERVE all facts: employers, roles, dates, bullets, tools, tech, certificates, languages, projects.
3) If the resume contains a section (Experience/Projects/etc.), it MUST appear in the output with its items (not empty).
4) You may rephrase for clarity, but keep the same items and chronology.
5) SKILLS is ONE STRING with TWO sublists:
   - "JD-Relevant (emphasis)" → put TECHNICAL items first that appear in BOTH the JD and the resume
     (programming languages, frameworks, DBs, cloud/devops, testing/tools). Soft skills only if the JD asks.
     Then append truly-missing items from MISSING SKILLS and mark with "(add)".
   - "All skills (from resume)" → verbatim list from the resume (dedupe case-insensitively).
6) CERTIFICATES & LANGUAGES: copy VERBATIM from the resume (single string with bullets). Only add if in MISSING SKILLS and mark with "(add)".
7) Output JSON uses **string content** for every section (no arrays/objects). Use bullets with "- " inside the string.
8) Never write “No professional experience listed …” or “No projects listed …”. If present in the resume, summarize the items.

Inputs:
---
RESUME TEXT (verbatim):
{use_text[:12000]}
---
JOB DESCRIPTION:
{job_description[:12000]}
---
MISSING SKILLS (minimal additions only, mark with "(add)"):
{', '.join(missing) if missing else 'none'}
---
USER HEADERS & CONTEXT (use these titles IN THIS ORDER and keep them all):
{header_text}

Output format — return ONLY valid JSON (no markdown/backticks) with string content in each section:
{{
  "sections": [
    {{"title":"Profile","content":"<polished but fact-identical summary>"}},
    {{"title":"Professional Experience","content":"<bulleted string; keep all employers/roles/dates/tools>"}},
    {{"title":"Education","content":"<bulleted string; unchanged facts>"}},
    {{"title":"Projects","content":"<bulleted string; keep all projects/tools/dates/impacts>"}},
    {{"title":"Certificates","content":"<bulleted string copied verbatim; optional '(add)'>"}},
    {{"title":"Skills","content":"JD-Relevant (emphasis):\\n- ...\\n- ...\\n\\nAll skills (from resume):\\n- ...\\n- ..."}},
    {{"title":"Languages","content":"<bulleted string copied verbatim with levels>"}}
  ]
}}
"""

    out = safe_invoke_chat(
        messages=[{"role":"user","content":[{"text":prompt}]}],
        system=[{"text":"Return ONLY valid JSON. No prose. No backticks."}],
        model_id=DEFAULT_CLAUDE,
        aws_region=DEFAULT_AWS_REGION,
        max_tokens=1400,
        temperature=0.2,
    )

    cleaned = out.strip().strip("`").replace("json", "", 1).strip()
    try:
        data = json.loads(cleaned)
    except Exception:
        data = {"raw": cleaned}

    # ----- normalize to string-only content + inject anchored fallbacks -----
    if isinstance(data, dict) and "sections" in data and isinstance(data["sections"], list):
        norm_sections = []
        for sec in data["sections"]:
            title = (sec.get("title") or "").strip()
            content = sec.get("content")

            # stringify any list/dict
            if isinstance(content, list):
                content = "\n".join(f"- {str(x)}" for x in content)
            elif isinstance(content, dict):
                lines = []
                for k, v in content.items():
                    if isinstance(v, list):
                        lines.append(f"{k}:")
                        lines += [f"- {str(x)}" for x in v]
                    else:
                        lines.append(f"- {k}: {v}")
                content = "\n".join(lines)
            elif content is None:
                content = ""

            # hard fallback: if model left it blank but we have an anchored block
            low = title.lower()
            if not str(content).strip():
                if low.startswith("professional experience") and exp_block:
                    content = exp_block
                if low.startswith("projects") and proj_block:
                    content = proj_block

            norm_sections.append({"title": title, "content": str(content)})
        data["sections"] = norm_sections
    else:
        if isinstance(data, dict):
            data = {"sections": [{"title": k, "content": str(v)} for k, v in data.items()]}

    return {"ok": True, "draft": data}
#------ CV docs draft build
def add_horizontal_line(paragraph):
    """Add a horizontal line under a paragraph"""
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2E5090')
    
    pBdr.append(bottom)
    pPr.append(pBdr)

def _format_skills_section(doc: Document, content: str):
    """
    Special formatting for skills section with:
    - JD-Relevant (emphasis) - highlighted
    - All skills (from resume) - regular
    """
    lines = [ln.strip() for ln in content.splitlines() if ln.strip()]
    
    current_category = None
    
    for line in lines:
        line_lower = line.lower()
        
        # Category headers
        if "jd-relevant" in line_lower or "emphasis" in line_lower:
            current_category = "relevant"
            p = doc.add_paragraph()
            p.space_before = Pt(4)
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(46, 80, 144)
            continue
        
        elif "all skills" in line_lower or "from resume" in line_lower:
            current_category = "all"
            p = doc.add_paragraph()
            p.space_before = Pt(8)
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(89, 89, 89)
            continue
        
        # Skill items
        if line.startswith("- "):
            p = doc.add_paragraph(style='List Bullet')
            text = line[2:].strip()
            
            # Check if it's a new skill (has "(add)" marker)
            if "(add)" in text.lower():
                text = text.replace("(add)", "").replace("(Add)", "").strip()
                run = p.add_run(text)
                run.font.color.rgb = RGBColor(0, 128, 0)  # Green for new skills
                
                # Add "(new)" badge
                badge = p.add_run(" (new)")
                badge.font.size = Pt(9)
                badge.font.italic = True
                badge.font.color.rgb = RGBColor(0, 128, 0)
            else:
                run = p.add_run(text)
                if current_category == "relevant":
                    run.bold = True
            
            p.space_after = Pt(2)
        else:
            # Inline comma-separated skills
            p = doc.add_paragraph(line)
            p.space_after = Pt(2)

def build_cv_docx_enhanced(
    full_name: str,
    contact: dict,
    sections: list,
) -> Document:
    """
    Build a professional, modern CV DOCX with:
    - Clean header with name and contact
    - Section headers with subtle accent lines
    - Professional font (Calibri/Arial fallback)
    - Proper spacing and margins
    - Bullet points for list items
    """
    doc = Document()
    
    # ---- Page Setup ----
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    
    # ---- Configure Normal Style ----
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
    
    # ---- HEADER: Name ----
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(full_name or "Your Name")
    name_run.font.name = 'Calibri'
    name_run.font.size = Pt(22)
    name_run.font.bold = True
    name_run.font.color.rgb = RGBColor(46, 80, 144)  # Professional blue
    
    # ---- HEADER: Contact Info ----
    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_para.space_after = Pt(12)
    
    contact_parts = []
    if contact.get('phone'):
        contact_parts.append(f"📱 {contact['phone']}")
    if contact.get('email'):
        contact_parts.append(f"✉ {contact['email']}")
    if contact.get('location'):
        contact_parts.append(f"📍 {contact['location']}")
    if contact.get('linkedin'):
        linkedin_clean = contact['linkedin'].replace('https://', '').replace('http://', '')
        contact_parts.append(f"🔗 {linkedin_clean}")
    
    contact_text = "  |  ".join(contact_parts)
    contact_run = contact_para.add_run(contact_text)
    contact_run.font.size = Pt(10)
    contact_run.font.color.rgb = RGBColor(89, 89, 89)
    
    # Add separator line after header
    separator = doc.add_paragraph()
    add_horizontal_line(separator)
    separator.space_after = Pt(8)
    
    # ---- SECTIONS ----
    for idx, sec in enumerate(sections):
        title = (sec.get("title") or "").strip()
        content = (sec.get("content") or "").strip()
        
        if not title and not content:
            continue
        
        # Section Title
        if title:
            section_para = doc.add_paragraph()
            section_para.space_before = Pt(12 if idx > 0 else 0)
            section_para.space_after = Pt(6)
            
            section_run = section_para.add_run(title.upper())
            section_run.font.name = 'Calibri'
            section_run.font.size = Pt(13)
            section_run.font.bold = True
            section_run.font.color.rgb = RGBColor(46, 80, 144)
            
            # Add subtle underline
            add_horizontal_line(section_para)
        
        if not content:
            continue
        
        # Process content
        lines = [ln for ln in content.splitlines() if ln.strip()]
        
        # Handle special sections
        title_lower = title.lower()
        
        if title_lower in ["profile", "summary", "objective"]:
            # Profile as regular paragraph
            para = doc.add_paragraph(content)
            para.space_after = Pt(6)
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
        elif title_lower == "skills":
            # Skills section with special formatting
            _format_skills_section(doc, content)
            
        else:
            # Default: process line by line
            for line in lines:
                stripped = line.strip()
                
                # Bullet point
                if stripped.startswith("- "):
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(stripped[2:].strip())
                    p.space_after = Pt(3)
                
                # Sub-bullet (double indent)
                elif stripped.startswith("  - ") or stripped.startswith("• "):
                    text = stripped[4:].strip() if stripped.startswith("  - ") else stripped[2:].strip()
                    p = doc.add_paragraph(style='List Bullet 2')
                    p.add_run(text)
                    p.space_after = Pt(2)
                    p.left_indent = Inches(0.5)
                
                # Bold-ish header lines
                elif re.match(r'^[A-Z].*\|.*$', stripped) or re.match(r'^[A-Z][^.]*$', stripped[:50]):
                    p = doc.add_paragraph()
                    p.space_before = Pt(6)
                    p.space_after = Pt(2)
                    run = p.add_run(stripped)
                    run.bold = True
                    run.font.size = Pt(11)
                
                # Regular text
                else:
                    p = doc.add_paragraph(stripped)
                    p.space_after = Pt(3)
        
        # Add spacing between sections
        if idx < len(sections) - 1:
            doc.add_paragraph()
    
    return doc
#------downl
@router.post("/cv_docx_download")
def cv_docx_download(
    request: Request,
    payload: Dict[str, Any] = Body(...),
):
    """
    Take the CV draft (sections, name, contact) from the frontend
    and return a nicely formatted DOCX file.
    Expected payload shape:
    {
      "full_name": "...",
      "contact": {
        "phone": "...",
        "email": "...",
        "location": "...",
        "linkedin": "..."
      },
      "sections": [
        {"title": "Profile", "content": "..."},
        ...
      ]
    }

    You can also pass the raw draft from /draft_cv_with_headers
    and just map it on the frontend to this shape.
    """
    # We’re not touching DB here, just using the token to ensure user is logged in
    # (optional: you can remove this if you want it completely open)
    # from db import get_db  # already imported at top
    # from auth import get_user_from_token  # already imported at top

    # Just to be safe, we won't crash if contact or sections are missing
    full_name = (payload.get("full_name") or "").strip() or "Your Name"

    contact_in = payload.get("contact") or {}
    contact = {
        "phone": contact_in.get("phone", ""),
        "email": contact_in.get("email", ""),
        "location": contact_in.get("location", ""),
        "linkedin": contact_in.get("linkedin", ""),
    }

    # sections can come directly, or from a nested "draft"
    sections = payload.get("sections")
    if sections is None and isinstance(payload.get("draft"), dict):
        sections = payload["draft"].get("sections")

    if not isinstance(sections, list):
        raise HTTPException(400, "sections must be a list of {title, content}.")

    # Normalize section items to avoid KeyError
    norm_sections: List[Dict[str, str]] = []
    for sec in sections:
        if not isinstance(sec, dict):
            continue
        norm_sections.append({
            "title": str(sec.get("title", "") or ""),
            "content": str(sec.get("content", "") or ""),
        })

    if not norm_sections:
        raise HTTPException(400, "No sections provided for DOCX export.")

    doc = build_cv_docx_enhanced(full_name, contact, norm_sections)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = (full_name.replace(" ", "_") or "HiringBuddy_CV") + ".docx"

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f'attachment; filename="{filename}"'
        },
    )

#--------interviewer agnet----


def _norm(s: Optional[str]) -> str:
    return (s or "").strip().lower()

# AUI program list 
AUI_PROGRAMS: Dict[str, Dict] = {
    "business administration": {
        "spec_label": "concentration",
        "options": [
            "marketing",
            "management",
            "finance",
            "international business",
            "logistics and supply chain management",
            "analytics and ai for business",
        ],
    },
    "international studies": {"spec_label": None, "options": []},
    "human resource development": {"spec_label": None, "options": []},
    "communication studies": {"spec_label": None, "options": []},
    "environmental studies and sustainability": {"spec_label": None, "options": []},
    "territorial planning and management": {
        "spec_label": "concentration",
        "options": ["environmental management", "applied geographic information systems"],
    },
    "psychology": {"spec_label": None, "options": []},

    "computer science": {
        "spec_label": "specialization",
        "options": [
            "advanced computer science",
            "artificial intelligence",
            "big data analytics",
            "computer systems",
            "software engineering",
        ],
    },
    "big data analytics": {"spec_label": None, "options": []},  # standalone BSc
    "cloud and mobile software design and development": {"spec_label": None, "options": []},  # standalone BSc

    "engineering and management science": {
        "spec_label": "thematic area",
        "options": [
            "manufacturing and logistics engineering",
            "decision support systems",
            "finance",
            "biotechnology",
        ],
    },
    "general engineering": {
        "spec_label": "thematic area",
        "options": ["mechatronics", "biotech", "aeronautics"],
    },
    "renewable energy systems engineering": {"spec_label": None, "options": []},
}

def _closest_option(major_norm: str, spec_norm: str) -> Optional[str]:
    """Loose match a user-typed specialization to the official options for their major, if any."""
    info = AUI_PROGRAMS.get(major_norm)
    if not info or not info["options"] or not spec_norm:
        return None
    for opt in info["options"]:
        o = _norm(opt)
        if spec_norm == o or spec_norm in o or o in spec_norm:
            return opt  # return official casing
    return None

@router.post("/interviewer/questions")
def interviewer_questions(
    request: Request,
    jd_text: str = Body(""),
    target_role: Optional[str] = Body(None),
    major: Optional[str] = Body(None),
    specialization: Optional[str] = Body(None),
    minor: Optional[str] = Body(None),
    language: Optional[str] = Body("en"), 
    db: Session = Depends(get_db),
):
    """
    Interviewer Agent (Claude-only):
    - Uses major + specialization + minor (if any).
    - Returns a *small* set of questions for cheap evaluation:
      2 intro + up to 5 technical + 1 behavioral (max 8).
    """
    lang_msg = _lang_hint(language) 
    user = get_user_from_token(request, db)
    if not user:
        raise HTTPException(status_code=401, detail="Unauthorized")

    prof = user.profile
    mj_raw = (major or (prof.major if prof else "") or "").strip()
    sp_raw = (specialization or (prof.specialization if prof else "") or "").strip()
    mn_raw = (minor or (prof.minor if prof else "") or "").strip()

    if not jd_text.strip():
        raise HTTPException(400, "Please provide job description text (jd_text).")
    if not mj_raw:
        raise HTTPException(400, "Profile major is missing. Add it in your profile or pass 'major' in the request.")

    mj_norm, sp_norm = _norm(mj_raw), _norm(sp_raw)
    prog = AUI_PROGRAMS.get(mj_norm, {"spec_label": None, "options": []})
    spec_label = prog["spec_label"]
    if sp_norm and prog["options"]:
        close = _closest_option(mj_norm, sp_norm)
        if close:
            sp_raw = close

    intro = [
        f"Present yourself. Mention your major{'' if not mn_raw else ' and minor'}, and highlight your AUI experience.",
        "What experience do you have that matches the requirements stated in the offer?",
        "What did you find interesting about this job offer?", "What is one project or experience you are proud of, and how does it relate to this role?",
    ]

    recognized = prog["options"]
    recognized_hint = (
        f"For this major, recognized AUI {spec_label or 'specializations'} include: {', '.join(recognized)}."
        if recognized else ""
    )
    spec_line = f"{(spec_label or 'specialization').title()}: {sp_raw}." if sp_raw else ""

    prompt = f"""
You are an interviewer generator for AUI students.
Generate concise, JD-aware interview questions tailored to the candidate.

Return ONLY valid JSON:
{{
  "technical": ["q1", "q2", "q3", "q4", "q5"],   // up to 5 single-sentence technical questions
  "behavioral": ["b1"]                           // 1 short behavioral question
}}

Candidate:
- Major: {mj_raw}.
- {spec_line}
- {('Minor: ' + mn_raw + '.') if mn_raw else ''}
- Target role: {target_role or 'unspecified'}.

AUI context:
{recognized_hint}

Guidelines:
- Prefer “explain/how would you/design” style over trivia.
- Use the JD’s tools/keywords where relevant.
- Avoid multiple questions in one; make each question stand alone.

Job Description (JD):
---
{jd_text[:12000]}
---
"""

    technical: List[str] = []
    behavioral: List[str] = []
    try:
        raw = safe_invoke_chat(
            messages=[{"role": "user", "content": [{"text": prompt}]}],
            system=[{"text": "Return ONLY valid JSON. No prose. No backticks."},
                    {"text": lang_msg},],
            model_id=DEFAULT_CLAUDE,
            aws_region=DEFAULT_AWS_REGION,
            max_tokens=350,
            temperature=0.3,
        )
        cleaned = raw.strip().strip("`").strip()
        m = re.search(r"\{.*\}", cleaned, re.DOTALL)
        json_str = m.group(0) if m else cleaned
        data = json.loads(json_str)

        if isinstance(data, dict):
            if isinstance(data.get("technical"), list):
                technical = [str(q).strip() for q in data["technical"] if isinstance(q, str)]
            if isinstance(data.get("behavioral"), list):
                behavioral = [str(q).strip() for q in data["behavioral"] if isinstance(q, str)]
    except Exception:
        technical = [
            "Walk me through how you would map the JD requirements to your coursework and projects.",
            "How would you design and validate a solution for the most critical requirement in this JD?",
        ]
        behavioral = [
            "Tell me about a time you had to quickly adapt to a new requirement or change of scope.",
        ]

    tech_qs = [q for q in technical if q][:5]   # 5tech
    beh_qs = [q for q in behavioral if q][:3]   #3 behavioral qts
    questions: List[Dict[str, str]] = []
    # ---------- PICK 2 RANDOM INTRO QUESTIONa----------
    random.shuffle(intro)
    selected_intros = intro[:2]

    for idx, q in enumerate(selected_intros):
        questions.append({
            "id": f"intro{idx+1}",
            "category": "intro",
            "text": q,
        })


    # # 1) Present yourself
    # questions.append({
    #     "id": "intro1",
    #     "category": "intro",
    #     "text": intro[0],
    # })

    # # 2) What experience matches the offer
    # if len(intro) > 1:
    #     questions.append({
    #         "id": "intro2",
    #         "category": "intro",
    #         "text": intro[1],
    #     })

    # # 3) What did you find interesting about the offer (optional third intro)
    # if len(intro) > 2:
    #     questions.append({
    #         "id": "intro3",
    #         "category": "intro",
    #         "text": intro[2],
    #     })

    # ---------- Technical questions from Claude ----------
    for i, q in enumerate(tech_qs):
        questions.append({
            "id": f"tech{i+1}",
            "category": "technical",
            "text": q,
        })

    # ---------- Behavioral questions from Claude ----------
    for i, b in enumerate(beh_qs):
        questions.append({
            "id": f"beh{i+1}",
            "category": "behavioral",
            "text": b,
        })

    # ---------- Hard cap at 8 questions total ----------
    
    questions = questions[:4]

    return {
        "ok": True,
        "major": mj_raw,
        "specialization": sp_raw,
        "minor_present": bool(mn_raw),
        "spec_label": spec_label,
        "recognized_specializations": recognized,
        "questions": questions,
    }


@router.post("/interviewer/evaluate")
def interviewer_evaluate(
    request: Request,
    jd_text: str = Body(...),
    major: Optional[str] = Body(None),
    specialization: Optional[str] = Body(None),
    answers: List[Dict[str, Any]] = Body(...),
    language: Optional[str] = Body("en"),
    db: Session = Depends(get_db),
):
    """
    Evaluate user's interview answers:
    - Scores each answer (0–5).
    - Returns ideal answer and feedback per question.
    - Computes final score (0–100) on the backend, regardless of what Claude returns.
    """
    user = get_user_from_token(request, db)
    if not user:
        raise HTTPException(status_code=401, detail="Unauthorized")

    if not jd_text.strip():
        raise HTTPException(400, "jd_text is required for context.")

    if not answers:
        raise HTTPException(400, "answers list is required and cannot be empty.")

    # Hard limit for testing / token saving
    answers = answers[:4]

    prof = user.profile
    mj_raw = (major or (prof.major if prof else "") or "").strip()
    sp_raw = (specialization or (prof.specialization if prof else "") or "").strip()

    # ---------- Build Q&A block ----------
    qa_block = []
    for item in answers:
        q = str(item.get("question", "")).strip()
        a = str(item.get("answer", "")).strip()
        if not q:
            continue
        # qa_block.append(f"Q: {q}\nA: {a or '[no answer]'}")
        short_a = (a or "").strip()
        if len(short_a) > 500:
            short_a = short_a[:500] + "..."
        qa_block.append(f"Q: {q}\nA: {short_a or '[no answer]'}")

    qa_text = "\n\n".join(qa_block)

    if not qa_text.strip():
        raise HTTPException(400, "No valid questions/answers supplied.")

    rubric = """
Score each answer from 0 to 5:
- 0 = no answer, completely off-topic, or incorrect.
- 1 = very weak, missing key ideas, shallow.
- 2 = partial but important gaps.
- 3 = acceptable, hits most core points.
- 4 = strong, clear, structured, uses relevant concepts/tools.
- 5 = excellent, specific, examples, metrics, and strong reasoning.

Weigh technical questions slightly more (they reflect role fit).
"""

    prompt = f"""
You are an interview coach for AUI students.

Candidate context:
- Major: {mj_raw or 'N/A'}
- Specialization: {sp_raw or 'N/A'}

Job Description:
---
{jd_text[:1200]}
---

The candidate answered these interview questions:
---
{qa_text[:12000]}
---

{rubric}

Your task:
For EACH question, produce:
- id: same id as the input question if you can infer it, otherwise "q1", "q2", ...
- score: integer 0–5
- ideal_answer: short model answer (4–7 sentences max)
- feedback: 1–2 sentences explaining what was good and what to improve.

Return ONLY valid JSON, no markdown, no backticks, with this structure:

{{
  "per_question": [
    {{
      "id": "q1",
      "score": 0,
      "ideal_answer": "string",
      "feedback": "string"
    }}
  ]
}}
"""

    raw = safe_invoke_chat(
        messages=[{"role": "user", "content": [{"text": prompt}]}],
        system=[{"text": "Return ONLY valid JSON. No prose. No backticks."},  {"text": _lang_hint(language)}, ],
        model_id=DEFAULT_CLAUDE,
        aws_region=DEFAULT_AWS_REGION,
        max_tokens=900,
        temperature=0.3,
    )

    print("\n=========== INTERVIEWER EVAL RAW ===========")
    print(raw)
    print("============================================\n")

    cleaned = raw.strip().strip("`").strip()
    m = re.search(r"\{.*\}", cleaned, re.DOTALL)
    json_str = m.group(0) if m else cleaned

    print("=========== INTERVIEWER EVAL CLEANED =======")
    print(json_str)
    print("============================================\n")

    # ---------- Parse JSON safely ----------
    try:
        data = json.loads(json_str)
    except Exception:
        data = {"per_question": []}

    # ---------- Normalize per-question ----------
    per_q_raw = data.get("per_question") or []
    if not isinstance(per_q_raw, list):
        per_q_raw = []

    # map id -> category from original answers
    id_to_cat: Dict[str, str] = {}
    for a in answers:
        aid = str(a.get("id") or "").strip()
        cat = str(a.get("category") or "").strip().lower()
        if aid:
            id_to_cat[aid] = cat

    norm_per_q = []
    for idx, item in enumerate(per_q_raw):
        if not isinstance(item, dict):
            continue

        # try to keep original ids so frontend matches
        input_id = answers[idx].get("id") if idx < len(answers) else None
        raw_id = item.get("id")
        fallback_id = f"q{idx+1}"
        qid = str(input_id or raw_id or fallback_id).strip()
        if input_id:
         qid = input_id
        score_raw = item.get("score", 0)
        try:
            if isinstance(score_raw, str):
                score_raw = score_raw.strip().replace("%", "")
            score_val = float(score_raw)
        except Exception:
            score_val = 0.0
        score_int = max(0, min(5, int(round(score_val))))

        norm_per_q.append({
            "id": qid,
            "score": score_int,
            "ideal_answer": str(item.get("ideal_answer", "")).strip(),
            "feedback": str(item.get("feedback", "")).strip(),
        })

    data["per_question"] = norm_per_q

  # ---------- ALWAYS compute final score here ----------
    total_weight = 0.0
    weighted_sum = 0.0

    for idx, q_eval in enumerate(norm_per_q):
        qid = q_eval["id"]
        score = q_eval["score"]
        cat = id_to_cat.get(qid, "")
        w = 1.5 if cat == "technical" else 1.0
        weighted_sum += score * w
        total_weight += w

    if total_weight > 0:
        raw_pct = (weighted_sum / (total_weight * 5.0)) * 100.0
        final_score = max(0, min(100, int(round(raw_pct))))
    else:
        final_score = 0

    strengths: List[str] = []
    improvements: List[str] = []
    resources: List[str] = []

    if norm_per_q:
        avg_score = sum(q["score"] for q in norm_per_q) / len(norm_per_q)
        if avg_score >= 4:
            strengths.append("Strong, clear technical communication and structured answers.")
        elif avg_score >= 3:
            strengths.append("Solid base of knowledge with generally relevant examples.")
        else:
            strengths.append("Basic understanding of concepts and willingness to learn.")

        low_questions = [q for q in norm_per_q if q["score"] <= 2]
        if low_questions:
            improvements.append("Some answers lack depth or concrete examples; add more details from your projects.")
            improvements.append("Clarify your reasoning steps instead of just listing tools.")
        else:
            improvements.append("Polish your examples with more metrics and impact where possible.")

        resources.append("Practice mock interviews focused on explaining your projects end-to-end.")
        resources.append("Review core backend concepts (REST APIs, databases, containerization).")
        resources.append("Record yourself answering common questions to improve clarity and confidence.")
    else:
        # fallback generic messages if Claude totally failed
        strengths.append("You started practicing interview questions, which is already a good step.")
        improvements.append("Try to give more detailed, structured answers using your projects as examples.")
        resources.append("Prepare 2–3 STAR stories from your projects and internships to reuse in different questions.")

    # ✅ define final_block OUTSIDE the if, so it always exists
    final_block = {
        "final_score": final_score,
        "strengths": strengths,
        "improvements": improvements,
        "resources": resources,
    }

    print("=========== INTERVIEWER FINAL BLOCK ========")
    print(final_block)
    print("============================================\n")

    # ---------- LOG INTERVIEW ATTEMPT----------
    try:
        job_title = (jd_text or "").strip().split("\n")[0][:255]
        jd_snippet = (jd_text or "")[:1000]

        attempt = InterviewAttempt(
            user_id=user.id,
            job_title=job_title,
            jd_snippet=jd_snippet,
            final_score=final_score,
            eval_json={
                "per_question": norm_per_q,
                "final": final_block,
            },
        )
        db.add(attempt)
        db.commit()
    except Exception as e:
        print("[WARN] Failed to log InterviewAttempt:", e)

    return {
        "ok": True,
        "evaluation": {
            "per_question": norm_per_q,
            "final": final_block,
        },
    }


#job search agent
@router.post("/job_search_serper")
def job_search_serper(
    request: Request,
    target_role: Optional[str] = Body(None),
    location: str = Body("Morocco"),
    num_results: int = Body(10),
    db: Session = Depends(get_db),
):
    """
    Use Serper.dev to find Moroccan job postings that fit the user's CV.
    - Reads latest resume for this user.
    - Uses Claude to extract skills + roles + seniority (job profile).
    - Builds a search query based on role, level, and skills.
    """
    if not SERPER_API_KEY:
        raise HTTPException(500, "SERPER_API_KEY not configured on the server.")

    user = get_user_from_token(request, db)
    if not user:
        raise HTTPException(status_code=401, detail="Unauthorized")

    # --- latest resume text ---
    resume = (
        db.query(Resume)
        .filter(Resume.user_id == user.id)
        .order_by(Resume.created_at.desc())
        .first()
    )
    if not resume or not (resume.text or "").strip():
        raise HTTPException(400, "No resume found. Upload & match a CV first.")

    full_text = resume.text or ""

    # --- AI profile from CV ---
    profile = {}
    skills = []
    try:
        profile = _extract_job_profile_from_resume(full_text)
        skills = profile.get("skills", []) or []
    except Exception as e:
        print("[WARN] _extract_job_profile_from_resume failed:", e)
        # skills = _extract_skill_tokens_from_resume(full_text, max_skills=10)
    if not skills:
        skills = _extract_skill_tokens_from_resume(full_text, max_skills=10)

    # keep profil
    if profile is None:
        profile = {}
    profile["skills"] = skills
    primary_roles = profile.get("primary_roles") or []
    seniority = (profile.get("seniority_level") or "student").lower()

    # --- decide base role  ---
    base_role = (target_role or "").strip()
    if not base_role and primary_roles:
        base_role = primary_roles[0]

    if not base_role:
        base_role = "internship" if seniority in {"student", "entry-level"} else "software engineer"

    level_hint = ""
    if seniority in {"student", "entry-level"}:
        level_hint = '(internship OR "entry level" OR junior)'
    elif seniority in {"junior"}:
        level_hint = '("junior" OR "1-3 years")'
    elif seniority in {"mid"}:
        level_hint = '("3-5 years" OR "mid-level")'
    elif seniority in {"senior"}:
        level_hint = '("senior" OR "5+ years")'

    # --- build Serper query ---
    skill_phrase = " ".join(skills[:8]) if skills else ""
    query = f'{base_role} {level_hint} jobs "{location}" {skill_phrase}'.strip()

    headers = {
        "X-API-KEY": SERPER_API_KEY,
        "Content-Type": "application/json",
    }
    payload = {
        "q": query,
        "num": num_results,
        "location": location,
        "gl": "ma",   # Morocco
        "hl": "fr",   # results mostly FR (you can switch to "en")
    }

    import httpx, re
    try:
        resp = httpx.post(SERPER_ENDPOINT, headers=headers, json=payload, timeout=10.0)
    except httpx.RequestError as e:
        raise HTTPException(502, f"Serper request error: {e}") from e

    if resp.status_code != 200:
        raise HTTPException(502, f"Serper error {resp.status_code}: {resp.text[:200]}")

    data = resp.json()
    organic = data.get("organic", []) or []

    jobish_domains = ("linkedin.com", "rekrute.com", "emploi.ma", "indeed.", "glassdoor.", "job", "careers")

    jobs = []
    for item in organic:
        title = item.get("title") or ""
        link = item.get("link") or ""
        snippet = item.get("snippet") or ""
        source = item.get("source") or ""

        text_for_domain = (link or source or "").lower()
        if not any(d in text_for_domain for d in jobish_domains):
        
            if not re.search(r"\b(emploi|job|stage|intern|engineer|developer|analyst|consultant)\b", title.lower()):
                continue

        # match CV skills against this posting
        combined = f"{title}\n{snippet}".lower()
        matched_skills = sorted(
            {s for s in skills if s.lower() in combined},
            key=lambda x: x.lower(),
        )

        jobs.append(
            {
                "title": title,
                "link": link,
                "snippet": snippet,
                "source": source,
                "position": item.get("position"),
                "matched_skills": matched_skills,
            }
        )

    return {
        "ok": True,
        "query": query,
        "location": location,
        "resume_id": resume.id,
        "skills_used": skills,
        "profile": profile,  
        "jobs": jobs[:num_results],
    }

#----------here for history and tarcking past macthes
@router.get("/match_history/recent")
def match_history_recent(
    request: Request,
    limit: int = 5,
    db: Session = Depends(get_db),
):
    user = get_user_from_token(request, db)
    if not user:
        raise HTTPException(status_code=401, detail="Unauthorized")

    rows = (
        db.query(MatchAttempt)
        .filter(MatchAttempt.user_id == user.id)
        .order_by(MatchAttempt.created_at.desc())
        .limit(limit)
        .all()
    )

    return {
        "items": [
            {
                "id": r.id,
                "job_title": r.job_title or "Unnamed job",
                "score": r.score,
                "resume_name": r.resume.name if r.resume else None,
                "created_at": r.created_at.isoformat() if r.created_at else None,
                "raw_json": r.raw_json,
            }
            for r in rows
        ]
    }


@router.get("/match_history")
def match_history_all(
    request: Request,
    db: Session = Depends(get_db),
):
    user = get_user_from_token(request, db)
    if not user:
        raise HTTPException(status_code=401, detail="Unauthorized")

    rows = (
        db.query(MatchAttempt)
        .filter(MatchAttempt.user_id == user.id)
        .order_by(MatchAttempt.created_at.desc())
        .all()
    )

    return {
        "items": [
            {
                "id": r.id,
                "job_title": r.job_title or "Unnamed job",
                "score": r.score,
                "resume_name": r.resume.name if r.resume else None,
                "created_at": r.created_at.isoformat() if r.created_at else None,
                "raw_json": r.raw_json,
            }
            for r in rows
        ]
    }
    # single match by id
@router.get("/match_history/{match_id}")
def match_history_one(
    request: Request,
    match_id: int,
    db: Session = Depends(get_db),
):
    user = get_user_from_token(request, db)
    if not user:
        raise HTTPException(status_code=401, detail="Unauthorized")

    r = (
        db.query(MatchAttempt)
        .filter(MatchAttempt.id == match_id, MatchAttempt.user_id == user.id)
        .first()
    )
    if not r:
        raise HTTPException(status_code=404, detail="Match not found")

    return {
        "id": r.id,
        "job_title": r.job_title or "Unnamed job",
        "score": r.score,
        "resume_name": r.resume.name if r.resume else None,
        "created_at": r.created_at.isoformat() if r.created_at else None,
        "raw_json": r.raw_json,
    }
# ---------- INTERVIEW HISTORY ---------

@router.get("/interview_history/recent")
def interview_history_recent(
    request: Request,
    limit: int = 5,
    db: Session = Depends(get_db),
):
    user = get_user_from_token(request, db)
    if not user:
        raise HTTPException(status_code=401, detail="Unauthorized")

    rows = (
        db.query(InterviewAttempt)
        .filter(InterviewAttempt.user_id == user.id)
        .order_by(InterviewAttempt.created_at.desc())
        .limit(limit)
        .all()
    )

    return {
        "items": [
            {
                "id": r.id,
                "job_title": r.job_title or "Untitled interview",
                "score": r.final_score,
                "created_at": r.created_at.isoformat() if r.created_at else None,
            }
            for r in rows
        ]
    }


@router.get("/interview_history")
def interview_history_all(
    request: Request,
    db: Session = Depends(get_db),
):
    user = get_user_from_token(request, db)
    if not user:
        raise HTTPException(status_code=401, detail="Unauthorized")

    rows = (
        db.query(InterviewAttempt)
        .filter(InterviewAttempt.user_id == user.id)
        .order_by(InterviewAttempt.created_at.desc())
        .all()
    )

    return {
        "items": [
            {
                "id": r.id,
                "job_title": r.job_title or "Untitled interview",
                "score": r.final_score,
                "created_at": r.created_at.isoformat() if r.created_at else None,
            }
            for r in rows
        ]
    }


@router.get("/interview_history/{interview_id}")
def interview_history_one(
    request: Request,
    interview_id: int,
    db: Session = Depends(get_db),
):
    user = get_user_from_token(request, db)
    if not user:
        raise HTTPException(status_code=401, detail="Unauthorized")

    r = (
        db.query(InterviewAttempt)
        .filter(InterviewAttempt.id == interview_id, InterviewAttempt.user_id == user.id)
        .first()
    )
    if not r:
        raise HTTPException(status_code=404, detail="Interview not found")

    return {
        "id": r.id,
        "job_title": r.job_title or "Untitled interview",
        "score": r.final_score,
        "created_at": r.created_at.isoformat() if r.created_at else None,
        "eval_json": r.eval_json,   # full blob: per_question + final
    }

# ---------------- memory clear no duplicate reset etc----------------
@router.post("/memory_clear")
def memory_clear(request: Request, db: Session = Depends(get_db)):
    user = get_user_from_token(request, db)
    db.query(ResumeChunk).filter(ResumeChunk.resume.has(user_id=user.id)).delete()
    db.query(Resume).filter(Resume.user_id == user.id).delete()
    db.commit()
    return {"ok": True}

@router.get("/memory_list")
def memory_list(request: Request, db: Session = Depends(get_db)):
    user = get_user_from_token(request, db)
    resumes = (
        db.query(Resume)
        .filter(Resume.user_id == user.id)
        .order_by(Resume.created_at.desc())
        .all()
    )
    return {"resumes": [{"id": r.id, "filename": r.name, "chars": len(r.text or "")} for r in resumes]}
